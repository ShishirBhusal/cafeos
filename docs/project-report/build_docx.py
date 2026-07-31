"""
CafeOS Project Report — Markdown to DOCX builder
================================================

Produces a Word document typeset to match the formatting of the reference
Tribhuvan University / D.A.V. College project report exactly.

Specification measured from the reference PDF:

  Page          US Letter, 8.5 x 11 in
  Margins       top 1.0", bottom 1.0", left 1.25", right 1.0"
  Body          Times New Roman 12 pt, justified, 1.5 line spacing (20.7 pt)
  Chapter head  Times New Roman 16 pt bold, centred, UPPERCASE, new page
  Section head  Times New Roman 14 pt bold, left
  Sub head      Times New Roman 12 pt bold, left
  Captions      Times New Roman 12 pt bold, centred
  Code          Courier New 10 pt
  Page numbers  centred footer, 11 pt; lower-roman front matter, arabic body
  Title page    TU logo centred at top, all text 12 pt centred

Run:  python build_docx.py
"""

from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except ImportError:
    Image = None

HERE = os.path.dirname(os.path.abspath(__file__))
REPORT_MD = os.path.join(HERE, "CafeOS_Project_Report.md")
FIGURES = os.path.join(HERE, "figures")
ML_REPORTS = os.path.normpath(os.path.join(HERE, "..", "..", "ml-service", "reports"))
OUT = os.path.join(HERE, "CafeOS_Project_Report.docx")

FONT = "Times New Roman"
MONO = "Courier New"
BODY_PT = 12
LINE_SPACING = 1.5
TEXT_WIDTH_IN = 6.25          # 8.5 - 1.25 - 1.0
MAX_IMG_W = 6.1
MAX_IMG_H = 7.4

# ---------------------------------------------------------------------------
# Title page content
# ---------------------------------------------------------------------------
TITLE = "CAFEOS: AN INTELLIGENT INVENTORY MANAGEMENT SYSTEM FOR CAFES"
STUDENT = "Rabindra Prasad Sah"
REG_NO = "[TU Reg No: ____________]"
COLLEGE = "D.A.V. College"
MONTH_YEAR = "July, 2026"
SUPERVISOR = "[Supervisor Name]"


# ===========================================================================
# Low-level helpers
# ===========================================================================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def add_field(paragraph, instr: str, size=11):
    """Insert a Word field code (used for PAGE numbers)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = FONT
    run.font.size = Pt(size)


def set_page_numbering(section, fmt: str, start: int | None = None):
    """fmt: 'decimal' | 'lowerRoman'"""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def add_footer_page_number(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    add_field(p, "PAGE")


def clear_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)


def style_para(p, size=BODY_PT, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0,
               space_before=0, line=LINE_SPACING):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    for r in p.runs:
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold or r.bold
        r.italic = italic or r.italic
    return p


# ===========================================================================
# Inline markdown
# ===========================================================================

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")


def add_inline(paragraph, text: str, size=BODY_PT, base_bold=False,
               bold=False, italic=False):
    """Render **bold**, *italic* and `code` inside a paragraph.

    Recurses so that markers nested inside one another -- most commonly a
    `code span` inside **bold** -- are rendered rather than leaking their
    literal backticks or asterisks into the document.
    """
    text = text.replace("<br>", " ").replace("&nbsp;", " ")
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            add_inline(paragraph, part[2:-2], size, base_bold, True, italic)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.font.name = MONO
            r.font.size = Pt(size - 1.5)
            r.bold = bold or base_bold
            r.italic = italic
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            add_inline(paragraph, part[1:-1], size, base_bold, bold, True)
        else:
            r = paragraph.add_run(part)
            r.font.name = FONT
            r.font.size = Pt(size)
            r.bold = bold or base_bold
            r.italic = italic


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*([^*\n]+?)\*", r"\1", text)
    text = re.sub(r"`([^`\n]+?)`", r"\1", text)
    text = text.replace("<br>", " ")
    return text.strip()


# ===========================================================================
# Images
# ===========================================================================

def image_size(path: str):
    if Image is None:
        return Inches(MAX_IMG_W), None
    with Image.open(path) as im:
        w, h = im.size
    ar = h / w
    width = MAX_IMG_W
    if width * ar > MAX_IMG_H:
        width = MAX_IMG_H / ar
    return Inches(width), None


def add_image(doc, path: str):
    if not os.path.exists(path):
        p = doc.add_paragraph()
        r = p.add_run(f"[missing image: {os.path.basename(path)}]")
        r.italic = True; r.font.name = FONT; r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    w, _ = image_size(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(path, width=w)


# ===========================================================================
# Title page
# ===========================================================================

def build_title_page(doc):
    def line(text, size=12, bold=False, italic=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        return p

    logo = os.path.join(FIGURES, "tu_logo.png")
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(logo, height=Inches(1.02))

    line("Tribhuvan University", 12)
    line("Faculty of Humanities and Social Sciences", 12, space_after=30)

    line(TITLE, 12, bold=True, space_before=18, space_after=30)
    line("A PROJECT REPORT", 12, bold=True, space_after=18)

    line("Submitted to", 12)
    line("Department of Computer Application", 12)
    line(COLLEGE, 12, space_after=48)

    line("In partial fulfillment of the requirements for the Bachelors in "
         "Computer Application", 12, italic=True, space_after=48)

    line("Submitted by", 12)
    line(STUDENT, 12)
    line(REG_NO, 12)
    line(MONTH_YEAR, 12, space_after=40)

    line("Under the Supervision of", 11)
    line(SUPERVISOR, 11)


# ===========================================================================
# Markdown parsing
# ===========================================================================

def parse_table(lines, i):
    """Consume a markdown table starting at line i. Returns (rows, next_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            txt = row[ci] if ci < len(row) else ""
            txt = txt.replace("<br>", " ")
            add_inline(p, txt, size=10.5, base_bold=(ri == 0))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.space_after = Pt(2)
            pf.space_before = Pt(2)

    # blank line after the table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    sp.paragraph_format.line_spacing = 1.0


CAPTION_RE = re.compile(r"^\*\*((?:Figure|Table)\s+[\d.]+:.*?)\*\*$")
IMG_PATH_RE = re.compile(r"^`([^`]+\.png)`$")


def build_body(doc, md: str, state: dict):
    lines = md.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.strip()

        # ---- skip artefacts -------------------------------------------
        if not line or line == "<br>" or line.startswith("<div style="):
            i += 1
            continue
        if line.startswith("---") and set(line) <= set("-"):
            i += 1
            continue

        # ---- fenced blocks --------------------------------------------
        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1

            if lang == "mermaid":
                state["mermaid"] += 1
                add_image(doc, os.path.join(FIGURES, f"diagram_{state['mermaid']:02d}.png"))
            else:
                for bl in block:
                    p = doc.add_paragraph()
                    r = p.add_run(bl if bl.strip() else " ")
                    r.font.name = MONO
                    r.font.size = Pt(9)
                    pf = p.paragraph_format
                    pf.line_spacing = 1.0
                    pf.space_after = Pt(0)
                    pf.left_indent = Inches(0.25)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(6)
                sp.paragraph_format.line_spacing = 1.0
            continue

        # ---- tables ----------------------------------------------------
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # ---- headings --------------------------------------------------
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), strip_md(m.group(2))
            if level == 1:
                if state["started"]:
                    doc.add_page_break()
                state["started"] = True
                p = doc.add_paragraph()
                r = p.add_run(text.upper())
                r.font.name = FONT; r.font.size = Pt(16); r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING
                pf.space_after = Pt(12)
            elif level == 2:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.name = FONT; r.font.size = Pt(14); r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING
                pf.space_before = Pt(10)
                pf.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.name = FONT; r.font.size = Pt(12); r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.line_spacing = LINE_SPACING
                pf.space_before = Pt(8)
                pf.space_after = Pt(2)
            i += 1
            continue

        # ---- figure / table caption ------------------------------------
        cap = CAPTION_RE.match(line)
        if cap:
            p = doc.add_paragraph()
            r = p.add_run(strip_md(cap.group(1)))
            r.font.name = FONT; r.font.size = Pt(12); r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            i += 1
            continue

        # ---- a bare image path (the Chapter 4 result figures) ----------
        img = IMG_PATH_RE.match(line)
        if img:
            rel = img.group(1)
            path = os.path.normpath(os.path.join(ML_REPORTS, "..", rel)) \
                if rel.startswith("reports/") else os.path.join(FIGURES, rel)
            if rel.startswith("reports/"):
                path = os.path.normpath(os.path.join(ML_REPORTS, rel[len("reports/"):]))
            add_image(doc, path)
            i += 1
            continue

        # ---- blockquote -------------------------------------------------
        if line.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            add_inline(p, " ".join(x for x in block if x), size=11)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.left_indent = Inches(0.35)
            pf.right_indent = Inches(0.2)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            continue

        # ---- bullets ----------------------------------------------------
        bm = re.match(r"^([-*])\s+(.*)$", line)
        if bm:
            indent = len(raw) - len(raw.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bm.group(2))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_after = Pt(0)
            pf.left_indent = Inches(0.3 + 0.25 * (indent // 2))
            i += 1
            continue

        # ---- checklist / numbered --------------------------------------
        nm = re.match(r"^(\d+)\.\s+(.*)$", line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, nm.group(2))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_after = Pt(0)
            pf.left_indent = Inches(0.35)
            i += 1
            continue

        # ---- ordinary paragraph -----------------------------------------
        block = [line]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "|", "```", ">", "- ", "* ", "<div"))
                    or re.match(r"^\d+\.\s", nxt) or CAPTION_RE.match(nxt)
                    or IMG_PATH_RE.match(nxt) or (nxt.startswith("---") and set(nxt) <= set("-"))):
                break
            block.append(nxt)
            i += 1

        p = doc.add_paragraph()
        add_inline(p, " ".join(block))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_after = Pt(6)


# ===========================================================================
# Assembly
# ===========================================================================

def main() -> int:
    with open(REPORT_MD, encoding="utf-8") as f:
        md = f.read()

    # Split off the markdown title block; the title page is built natively.
    marker = "## ACKNOWLEDGEMENT"
    if marker not in md:
        print("ERROR: could not locate '## ACKNOWLEDGEMENT'")
        return 1
    rest = md[md.index(marker):]

    # Front matter ends where Chapter 1 begins.
    ch1 = "# CHAPTER 1: INTRODUCTION"
    if ch1 not in rest:
        print("ERROR: could not locate Chapter 1")
        return 1
    front, body = rest[:rest.index(ch1)], rest[rest.index(ch1):]

    doc = Document()

    # ---- base style ----
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(0)

    # ---- SECTION 1: title page (no page number) ----
    s1 = doc.sections[0]
    s1.page_width, s1.page_height = Inches(8.5), Inches(11)
    s1.top_margin = s1.bottom_margin = s1.right_margin = Inches(1.0)
    s1.left_margin = Inches(1.25)
    clear_footer(s1)
    build_title_page(doc)

    # ---- SECTION 2: front matter (lower roman, starting ii) ----
    s2 = doc.add_section(WD_SECTION.NEW_PAGE)
    s2.page_width, s2.page_height = Inches(8.5), Inches(11)
    s2.top_margin = s2.bottom_margin = s2.right_margin = Inches(1.0)
    s2.left_margin = Inches(1.25)
    set_page_numbering(s2, "lowerRoman", start=2)
    add_footer_page_number(s2)

    state = {"mermaid": 0, "started": False}
    build_body(doc, front, state)

    # ---- SECTION 3: body (arabic, starting 1) ----
    s3 = doc.add_section(WD_SECTION.NEW_PAGE)
    s3.page_width, s3.page_height = Inches(8.5), Inches(11)
    s3.top_margin = s3.bottom_margin = s3.right_margin = Inches(1.0)
    s3.left_margin = Inches(1.25)
    set_page_numbering(s3, "decimal", start=1)
    add_footer_page_number(s3)

    state["started"] = False
    build_body(doc, body, state)

    doc.save(OUT)

    size_kb = os.path.getsize(OUT) / 1024
    print("=" * 66)
    print("CafeOS Project Report — DOCX built")
    print("=" * 66)
    print(f"  Output       : {OUT}")
    print(f"  Size         : {size_kb:,.0f} KB")
    print(f"  Diagrams     : {state['mermaid']} mermaid figures embedded")
    print(f"  Font         : {FONT} {BODY_PT} pt, {LINE_SPACING} line spacing")
    print(f"  Page         : US Letter, margins L1.25\" T/B/R 1.0\"")
    print(f"  Numbering    : lower-roman front matter, arabic from Chapter 1")
    print("=" * 66)
    return 0


if __name__ == "__main__":
    sys.exit(main())
